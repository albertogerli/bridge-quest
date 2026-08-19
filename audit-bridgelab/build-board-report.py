from __future__ import annotations

import math
import os
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path("/Users/albertogiovannigerli/Desktop/Personale/Bridge/bridgequest")
OUTPUT = ROOT / "perizia" / "Relazione_BridgeLab_per_Consiglio_Federale_2026-08-15.docx"
ASSET_DIR = Path("/tmp/bridgelab-board-report-assets")
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

SKILL_SCRIPTS = Path(
    "/Users/albertogiovannigerli/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.813.12317/skills/documents/scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


# standard_business_brief token map, with a named report-cover override.
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)
HEADER_DISTANCE = Inches(0.492)
FOOTER_DISTANCE = Inches(0.492)
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
FONT = "Calibri"
NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D0D5DD"
WHITE = "FFFFFF"
GOLD = "B7791F"
GREEN = "276749"
LIGHT_GREEN = "E8F5EE"
RED = "9B1C1C"
LIGHT_RED = "FDECEC"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, name=FONT, size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in edge_data:
                element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_cell_text(cell, text, *, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    set_run_font(p.add_run(str(text)), size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, widths, header=True, font_size=9.5, numeric_cols=()):
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )
    table.style = "Table Grid"
    if header:
        set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        prevent_row_split(row)
        for col_idx, cell in enumerate(row.cells):
            if row_idx == 0 and header:
                shade_cell(cell, LIGHT_GRAY)
                for run in cell.paragraphs[0].runs:
                    set_run_font(run, size=font_size, color=NAVY, bold=True)
            else:
                for run in cell.paragraphs[0].runs:
                    set_run_font(run, size=font_size, color=INK)
            if col_idx in numeric_cols:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, weights, *, numeric_cols=(), font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], value, bold=True, color=NAVY, size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.RIGHT if i in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, size=font_size, align=align)
    style_table(table, column_widths_from_weights(weights, CONTENT_WIDTH_DXA), font_size=font_size, numeric_cols=numeric_cols)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_table_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run(text), size=8.5, color=MUTED, italic=True)
    return p


def add_numbering_definition(doc, marker="•"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), marker)
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, text, num_id, *, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), size=11, bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]), size=11)
    else:
        set_run_font(p.add_run(text), size=11)
    return p


def add_body(doc, text, *, bold_lead=None, italic=False, color=INK, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_together = keep
    p.paragraph_format.keep_with_next = keep
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), size=11, color=color, bold=True)
        set_run_font(p.add_run(text[len(bold_lead):]), size=11, color=color, italic=italic)
    else:
        set_run_font(p.add_run(text), size=11, color=color, italic=italic)
    return p


def add_callout(doc, label, text, *, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": accent})
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    set_run_font(p.add_run(label + "  "), size=10.5, color=accent, bold=True)
    set_run_font(p.add_run(text), size=10.5, color=INK)
    apply_table_geometry(
        table,
        [CONTENT_WIDTH_DXA],
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=180,
        cell_margins_dxa={"top": 140, "bottom": 140, "start": 180, "end": 180},
    )
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_section_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run(text), size=8.5, color=MUTED, italic=True)


def add_picture_with_alt(doc, path, *, width, alt_text):
    shape = doc.add_picture(str(path), width=width)
    shape._inline.docPr.set("title", alt_text)
    shape._inline.docPr.set("descr", alt_text)
    return shape


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=9, color=MUTED)


FONT_REG = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"


def pil_font(size, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REG, size)


def draw_bar_chart(path, labels, values, title, *, max_value=None, color=(46, 116, 181), value_suffix=""):
    width, height = 1500, 760
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 42), title, font=pil_font(38, True), fill=(23, 54, 93))
    left, top, right, bottom = 120, 130, 1430, 650
    d.line((left, bottom, right, bottom), fill=(170, 178, 189), width=2)
    ceiling = max_value or max(values) * 1.12
    slot = (right - left) / len(values)
    bar_w = slot * 0.58
    for i, (label, value) in enumerate(zip(labels, values)):
        x0 = left + i * slot + (slot - bar_w) / 2
        x1 = x0 + bar_w
        y0 = bottom - (value / ceiling) * (bottom - top)
        d.rounded_rectangle((x0, y0, x1, bottom), radius=10, fill=color)
        value_text = f"{value:g}{value_suffix}"
        box = d.textbbox((0, 0), value_text, font=pil_font(24, True))
        d.text(((x0 + x1 - (box[2]-box[0])) / 2, y0 - 34), value_text, font=pil_font(24, True), fill=(31, 41, 55))
        lbox = d.textbbox((0, 0), label, font=pil_font(21))
        d.text(((x0 + x1 - (lbox[2]-lbox[0])) / 2, bottom + 18), label, font=pil_font(21), fill=(71, 84, 103))
    img.save(path, dpi=(180, 180))


def draw_work_pattern(path):
    width, height = 1800, 840
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 38), "Distribuzione dei commit: quando è stato svolto il lavoro", font=pil_font(40, True), fill=(23, 54, 93))
    panels = [
        ((70, 130, 870, 735), ["Lun", "Mar", "Mer", "Gio", "Ven", "Sab", "Dom"], [32, 57, 60, 56, 38, 36, 50], "Per giorno della settimana", (46, 116, 181)),
        ((940, 130, 1730, 735), ["00–05", "06–11", "12–17", "18–23"], [58, 66, 102, 103], "Per fascia oraria", (183, 121, 31)),
    ]
    for (left, top, right, bottom), labels, values, subtitle, color in panels:
        d.text((left, top), subtitle, font=pil_font(28, True), fill=(31, 77, 120))
        chart_top = top + 85
        chart_bottom = bottom - 55
        d.line((left + 40, chart_bottom, right - 20, chart_bottom), fill=(170, 178, 189), width=2)
        slot = (right - left - 80) / len(values)
        ceiling = max(values) * 1.15
        for i, (label, value) in enumerate(zip(labels, values)):
            x0 = left + 50 + i * slot
            x1 = x0 + slot * .55
            y0 = chart_bottom - value / ceiling * (chart_bottom - chart_top)
            d.rounded_rectangle((x0, y0, x1, chart_bottom), radius=8, fill=color)
            val = str(value)
            vb = d.textbbox((0, 0), val, font=pil_font(22, True))
            d.text(((x0+x1-vb[2]+vb[0])/2, y0-30), val, font=pil_font(22, True), fill=(31, 41, 55))
            lb = d.textbbox((0, 0), label, font=pil_font(19))
            d.text(((x0+x1-lb[2]+lb[0])/2, chart_bottom+14), label, font=pil_font(19), fill=(71, 84, 103))
    img.save(path, dpi=(200, 200))


def draw_activity_line(path):
    values = [71, 67, 67, 59, 57, 69, 63, 61, 72, 66, 66, 72, 71, 79]
    labels = ["02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12", "13", "14", "15"]
    width, height = 1600, 690
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 38), "Utenti attivi giornalieri — ultimi 14 giorni", font=pil_font(38, True), fill=(23, 54, 93))
    left, top, right, bottom = 120, 145, 1510, 565
    min_y, max_y = 50, 85
    for tick in [50, 60, 70, 80]:
        y = bottom - (tick-min_y)/(max_y-min_y)*(bottom-top)
        d.line((left, y, right, y), fill=(230, 233, 238), width=2)
        d.text((65, y-12), str(tick), font=pil_font(18), fill=(102, 112, 133))
    pts = []
    for i, value in enumerate(values):
        x = left + i * (right-left)/(len(values)-1)
        y = bottom - (value-min_y)/(max_y-min_y)*(bottom-top)
        pts.append((x, y))
    d.line(pts, fill=(46, 116, 181), width=7, joint="curve")
    for i, ((x, y), value, label) in enumerate(zip(pts, values, labels)):
        d.ellipse((x-8, y-8, x+8, y+8), fill=(183, 121, 31), outline="white", width=3)
        if i in [0, len(values)-1] or value in [min(values), max(values)]:
            d.text((x-12, y-38), str(value), font=pil_font(18, True), fill=(31, 41, 55))
        d.text((x-10, bottom+18), label, font=pil_font(16), fill=(102, 112, 133))
    d.text((120, 620), "Media 67,14  |  minimo 57  |  massimo 79", font=pil_font(22, True), fill=(39, 103, 73))
    img.save(path, dpi=(190, 190))


def draw_value_scenarios(path):
    scenarios = [
        ("Basso", 78.41613, 2.58407, 81.00021),
        ("Centrale", 115.18571, 101.08272, 216.26843),
        ("Alto", 154.05175, 296.54984, 450.60159),
    ]
    width, height = 1550, 760
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 40), "Valore attuale del contratto — scenari", font=pil_font(40, True), fill=(23, 54, 93))
    d.text((70, 93), "Migliaia di euro; quota iniziale all-inclusive + valore residuo della licenza", font=pil_font(22), fill=(102, 112, 133))
    left, top, right, bottom = 170, 170, 1440, 650
    ceiling = 500
    slot = (right-left)/3
    bar_w = 170
    for i, (label, initial, residual, total) in enumerate(scenarios):
        x0 = left + i*slot + (slot-bar_w)/2
        x1 = x0 + bar_w
        initial_h = initial/ceiling*(bottom-top)
        residual_h = residual/ceiling*(bottom-top)
        y_mid = bottom-initial_h
        y_top = y_mid-residual_h
        d.rounded_rectangle((x0, y_mid, x1, bottom), radius=10, fill=(46,116,181))
        if residual > 0:
            d.rounded_rectangle((x0, y_top, x1, y_mid+10), radius=10, fill=(183,121,31))
        total_text = f"€{total:.0f}k"
        tb = d.textbbox((0,0), total_text, font=pil_font(27, True))
        d.text(((x0+x1-tb[2]+tb[0])/2, y_top-40), total_text, font=pil_font(27, True), fill=(31,41,55))
        lb = d.textbbox((0,0), label, font=pil_font(23, True))
        d.text(((x0+x1-lb[2]+lb[0])/2, bottom+18), label, font=pil_font(23, True), fill=(31,77,120))
    d.rectangle((190, 698, 222, 730), fill=(46,116,181))
    d.text((235, 699), "Primi 3 anni all-inclusive", font=pil_font(20), fill=(71,84,103))
    d.rectangle((700, 698, 732, 730), fill=(183,121,31))
    d.text((745, 699), "Licenza residua", font=pil_font(20), fill=(71,84,103))
    img.save(path, dpi=(190, 190))


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.right_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.header_distance = HEADER_DISTANCE
    section.footer_distance = FOOTER_DISTANCE

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, color, before, after in [
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_run_font(hp.add_run("BRIDGELAB  |  RELAZIONE PER IL CONSIGLIO FEDERALE"), size=8.5, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    set_run_font(fp.add_run("15 agosto 2026   •   "), size=8.5, color=MUTED)
    add_page_number(fp)


def add_cover(doc):
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    set_run_font(p.add_run("RELAZIONE INDIPENDENTE"), size=11, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run("BridgeLab"), size=31, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    set_run_font(p.add_run("Funzionalità, adozione, effort di sviluppo, maturità infrastrutturale e valore della licenza"), size=15, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(52)
    set_run_font(p.add_run("Destinata al Consiglio Federale\nRilevazione al 15 agosto 2026"), size=11, color=MUTED, italic=True)

    table = doc.add_table(rows=1, cols=3)
    metrics = [("1.095", "utenti registrati"), ("8.221 h 48 m", "tempo totale in app"), ("€216 mila", "valore centrale stimato")]
    for idx, (value, label) in enumerate(metrics):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(value), size=18, color=NAVY, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        set_run_font(p2.add_run(label), size=9, color=MUTED)
    apply_table_geometry(table, [3120, 3120, 3120], table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=TABLE_INDENT_DXA, cell_margins_dxa={"top": 180, "bottom": 180, "start": 120, "end": 120})
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Documento privo di dati personali • valori misurati, calcolati o stimati distintamente"), size=9, color=MUTED)
    doc.add_page_break()


def build_document():
    work_chart = ASSET_DIR / "git_work_patterns.png"
    activity_chart = ASSET_DIR / "daily_activity.png"
    value_chart = ASSET_DIR / "licence_value.png"
    draw_work_pattern(work_chart)
    draw_activity_line(activity_chart)
    draw_value_scenarios(value_chart)

    doc = Document()
    configure_styles(doc)
    bullet_id = add_numbering_definition(doc)
    add_cover(doc)

    add_section_heading(doc, "Sintesi esecutiva", 1)
    add_callout(
        doc,
        "CONCLUSIONE",
        "BridgeLab è una piattaforma verticale già operativa, con contenuti, gioco, strumenti per istruttori e infrastruttura di livello produttivo. Il valore attuale centrale del pacchetto contrattuale — tre anni all-inclusive più licenza italiana a tempo indeterminato — è stimato in €216.268, con un intervallo prudenziale di €81.000–€450.602.",
        fill=LIGHT_GREEN,
        accent=GREEN,
    )
    add_body(doc, "Il prodotto non va valutato soltanto per quanto costerebbe riscriverlo. Il beneficio per la Federazione deriva dalla disponibilità immediata di una piattaforma specialistica, dall'uso già osservato, dall'esclusiva italiana per finalità istituzionali e didattiche, dai servizi gratuiti fino al 31 dicembre 2029 e dalla continuità sul codice prevista dal contratto.")
    add_table(
        doc,
        ["Indicatore", "Evidenza aggiornata", "Lettura direzionale"],
        [
            ("Prodotto", "81 pagine, 102 componenti, 22 famiglie di flusso", "Piattaforma completa, non dimostrativa"),
            ("Contenuti", "4 corsi, 49 lezioni, 199 moduli vivi", "Catalogo didattico già utilizzabile"),
            ("Adozione", "1.095 utenti; 60.773 partite", "Uso reale e ripetuto"),
            ("Tempo in app", "8.221 h 48 m; media 450 minuti/utente", "Profondità d'uso, non sola registrazione"),
            ("Qualità tecnica", "3,42/5; 990 test superati", "Buona maturità con rischi ancora gestibili"),
            ("Effort", "329 commit; 68,09% fuori dalla finestra feriale 09–18", "Sviluppo fortemente concentrato ed extra-orario"),
            ("Valore contratto", "€216.268 centrale; €81.000–€450.602", "Valore d'uso, non costo di sostituzione"),
        ],
        [1.35, 2.35, 2.80],
        font_size=9.2,
    )
    add_table_note(doc, "Fonti: repository al commit 477d2cd6; dashboard amministrativa aggiornata alle 21:14:18; contratto di licenza in bozza; calcoli riproducibili in appendice.")

    add_section_heading(doc, "1. Analisi completa delle funzionalità", 1)
    add_body(doc, "BridgeLab riunisce in un unico ambiente ciò che, sul mercato, è spesso distribuito tra corso online, motore di gioco, strumenti per l'istruttore e comunità. La completezza non è omogenea: il nucleo didattico e di gioco ha uso consolidato; alcune funzioni collaborative introdotte di recente sono tecnicamente complete ma ancora poco utilizzate.")

    add_section_heading(doc, "1.1 Apprendimento strutturato", 2)
    for text in [
        "Quattro corsi organizzati in 16 mondi, con 49 lezioni e 199 moduli nel catalogo operativo.",
        "Contenuti teorici, esempi, regole, esercizi e cinque tipologie di quiz integrate nelle lezioni.",
        "111 domande di comprensione, 31 esercizi di eserciziario, 32 scenari “Trova l'errore”, 32 scenari d'impasse e 20 scenari di pratica dichiarativa.",
        "Glossario di 49 voci, progressione per prerequisiti, completamenti, ripasso degli errori e ripetizione dilazionata.",
    ]:
        add_bullet(doc, text, bullet_id)

    add_section_heading(doc, "1.2 Gioco, allenamento e analisi", 2)
    for text in [
        "Gioco completo della smazzata con dichiarazione, carta, dummy, prese, punteggio, risultato e analisi finale.",
        "272 mani di catalogo con analisi double-dummy e 1.504 mani generate/condivise con distribuzioni di valore atteso.",
        "Modalità di pratica: smazzata, sfida, mano del giorno, conta veloce, quiz lampo, memory, impasse, dichiarazione, mano guidata, pratica licita e compiti.",
        "Confronto sullo stesso board, sfide tra amici, sfide a coppie, tornei di licita e classifiche.",
        "Integrazione con DDS per l'analisi esatta e con BEN per suggerimenti e gioco automatico, con timeout e fallback.",
    ]:
        add_bullet(doc, text, bullet_id)

    add_section_heading(doc, "1.3 Scuola, istruttori e comunità", 2)
    add_table(
        doc,
        ["Area", "Funzioni disponibili", "Evidenza d'uso"],
        [
            ("Portale istruttori", "Classi, compiti, chat, classifiche, archivio, studio mani, lavagna e dispense", "16 istruttori, 16 classi, 52 allievi"),
            ("Generazione mani", "Vincoli, double-dummy, par, PBN e assegnazione alla classe", "Implementata; uso non isolato nell'allegato"),
            ("Tavolo live", "Posti, carte nascoste, gioco condiviso, reveal e fallback polling", "Uso iniziale: 1 tavolo nella rilevazione tecnica"),
            ("Comunità", "Amici, forum, circoli, ricerca compagno, sfide e leaderboard", "Uso osservato sui flussi storici"),
            ("Amministrazione", "Dashboard, consensi, segmentazione, richieste istruttore, account e comunicazioni", "Dashboard operativa aggiornata"),
        ],
        [1.25, 3.25, 2.00],
        font_size=9.1,
    )

    add_section_heading(doc, "1.4 Completezza effettiva", 2)
    add_table(
        doc,
        ["Stato", "Numero di famiglie", "Interpretazione"],
        [
            ("Funzionanti con uso osservato", "11", "Nucleo didattico, gioco, socialità, classi e ripasso"),
            ("Implementate, uso non misurato separatamente", "6", "Percorso completo nel codice, metrica specifica assente"),
            ("Con uso iniziale", "3", "Tavolo live, campo condiviso e torneo di licita"),
            ("Implementate senza uso osservato", "2", "Licita a due e sfida a coppie"),
        ],
        [2.35, 1.15, 3.00],
        numeric_cols=(1,),
    )
    add_callout(doc, "LETTURA", "La piattaforma è completa nel nucleo che genera oggi la maggior parte dell'utilizzo. Le funzioni più recenti rappresentano valore opzionale e potenziale, non ancora pienamente dimostrato dall'adozione.")

    add_section_heading(doc, "2. Effort di sviluppo e lavoro extra-orario", 1)
    add_body(doc, "La storia Git è la fonte più oggettiva disponibile per descrivere continuità e intensità del lavoro. Non equivale a un timesheet: un commit registra un risultato, ma non le ore di progettazione, analisi, debug o preparazione che lo precedono.")
    add_table(
        doc,
        ["Metrica", "Risultato", "Significato"],
        [
            ("Commit complessivi", "329", "328 attribuiti allo sviluppatore principale"),
            ("Periodo", "8 febbraio–15 agosto 2026", "Sviluppo concentrato in poco più di sei mesi"),
            ("Giorni con almeno un commit", "57", "Presenza distribuita nel calendario"),
            ("Sessioni ricostruite", "102", "Nuova sessione dopo oltre 90 minuti di intervallo"),
            ("Ore ricostruite", "127,84", "Limite inferiore, non effort totale"),
            ("Dimensione mediana commit", "168 righe", "Indicatore di ampiezza media delle consegne"),
        ],
        [2.05, 1.55, 2.90],
        numeric_cols=(1,),
    )
    add_picture_with_alt(
        doc,
        work_chart,
        width=Inches(6.35),
        alt_text="Grafico a barre dei commit per giorno della settimana, suddivisi nelle quattro fasce orarie 00–05, 06–11, 12–17 e 18–23.",
    )
    add_caption(doc, "Figura 1 — Commit per giorno e fascia oraria, fuso Europe/Rome.")
    doc.add_page_break()

    add_table(
        doc,
        ["Indicatore extra-orario", "Commit", "Quota"],
        [
            ("Sabato o domenica", "86", "26,14%"),
            ("Dalle 18:00 alle 05:59", "161", "48,94%"),
            ("Weekend o feriale fuori 09:00–17:59", "224", "68,09%"),
        ],
        [3.55, 1.15, 1.80],
        numeric_cols=(1, 2),
    )
    add_body(doc, "I 57 giorni attivi comprendono 18 giorni di weekend e 39 feriali. Domenica conta 50 commit e sabato 36; la fascia 18:00–23:59 è la più popolata con 103 commit, seguita dalla fascia 12:00–17:59 con 102.")
    add_callout(doc, "INTERPRETAZIONE PRUDENTE", "Il 68,09% è un proxy di attività fuori da una finestra lavorativa convenzionale, non prova ore contrattuali né misura la fatica personale. Dimostra però che una quota maggioritaria delle consegne è avvenuta di sera, di notte, al mattino presto o nel fine settimana.", fill="FFF8E8", accent=GOLD)
    add_table(
        doc,
        ["Mese 2026", "Feb", "Mar", "Apr", "Mag", "Giu", "Lug", "Ago*"],
        [("Commit", "53", "96", "7", "21", "20", "14", "118")],
        [1.35, .74, .74, .74, .74, .74, .74, .71],
        numeric_cols=(1, 2, 3, 4, 5, 6, 7),
        font_size=9.1,
    )
    add_table_note(doc, "* Agosto fino al 15. I 118 commit evidenziano una fase finale particolarmente intensa.")

    add_section_heading(doc, "3. Crescita, utenti e profondità d'uso", 1)
    add_body(doc, "La dashboard amministrativa aggiornata alle 21:14:18 mostra che l'utilizzo non si limita alle registrazioni: sono presenti tempo accumulato, attività quotidiana, partite e una base di istruttori e classi. Tutti i dati riportati sono aggregati; nessuna informazione personale è stata trasferita nella relazione.")
    add_table(
        doc,
        ["Indicatore", "Valore", "Lettura"],
        [
            ("Utenti registrati", "1.095", "Base complessiva"),
            ("Nuovi utenti", "3 oggi • 12 in 7 gg • 49 in 30 gg", "+4,68% rispetto alla base iniziale dei 30 giorni"),
            ("Utenti attivi", "79 oggi • 133 in 7 gg", "7,21% e 12,15% della base registrata"),
            ("Attività giornaliera, 14 gg", "media 67,14 • minimo 57 • massimo 79", "Andamento stabile nel periodo breve"),
            ("Tempo complessivo", "8.221 h 48 m", "493.308 minuti aggregati"),
            ("Tempo medio", "450 minuti per utente", "Circa 7 ore e 30 minuti per registrato"),
            ("Retention a 7 giorni", "11%", "Metrica operativa della dashboard"),
        ],
        [2.05, 2.35, 2.10],
        font_size=9.2,
    )
    add_picture_with_alt(
        doc,
        activity_chart,
        width=Inches(6.35),
        alt_text="Grafico lineare degli utenti attivi giornalieri negli ultimi quattordici giorni, da un minimo di 57 a un massimo di 79.",
    )
    add_caption(doc, "Figura 2 — Attivi giornalieri: media 67,14, intervallo 57–79.")

    add_section_heading(doc, "3.1 Attività di gioco", 2)
    add_table(
        doc,
        ["Metrica", "Valore"],
        [
            ("Partite complessive", "60.773"),
            ("Mani giocate, indicatore dashboard", "38.891"),
            ("Partite negli ultimi 7 giorni", "2.566"),
            ("Partite nel giorno", "389"),
            ("Giocatori unici complessivi", "665"),
            ("Giocatori negli ultimi 7 giorni", "106"),
            ("Quota smazzate + sfide", "74,11% delle partite"),
            ("XP complessivo / medio", "7.345.518 / 6.708"),
            ("Streak massimo", "98"),
        ],
        [3.85, 2.65],
        numeric_cols=(1,),
    )
    add_body(doc, "Le due modalità principali sono le smazzate complete, con 34.062 risultati, e le sfide fra utenti, con 10.978. Seguono mano del giorno (3.591), conta veloce (2.583), quiz lampo (1.913) e memory (1.876). La somma delle 13 categorie riportate dalla dashboard coincide con 60.773 partite.")

    add_section_heading(doc, "3.2 Profilo e canali di accesso", 2)
    add_body(doc, "Il pubblico registrato è composto per il 64% da profili senior, per il 31% da adulti e per il 5% da giovani; i profili junior risultano pari a zero nella rilevazione. Negli ultimi 30 giorni web e PWA rappresentano insieme il 90% degli accessi, iOS il 9% e Android lo 0%. Questi dati confermano l'adeguatezza del canale web/PWA, ma indicano uno spazio di miglioramento sul mobile nativo e sul reclutamento giovanile.")

    add_section_heading(doc, "4. Punti di forza rispetto ai concorrenti", 1)
    add_body(doc, "Il confronto non attribuisce a BridgeLab una superiorità assoluta. Funbridge e BBO dispongono di comunità internazionali, intelligenze di gioco mature e ampia disponibilità multipiattaforma. Il vantaggio di BridgeLab è diverso: integrazione verticale con la didattica federale, controllo istituzionale e capacità di modificare il prodotto sul processo italiano.")
    add_table(
        doc,
        ["Dimensione", "BridgeLab", "Funbridge / BBO", "LMS generico"],
        [
            ("Curriculum", "Percorsi FIGB, lezioni, quiz, mani e progressione nello stesso ambiente", "Didattica e pratica internazionale; non costruite sul processo FIGB", "Forte gestione corsi, senza logica bridge nativa"),
            ("Gioco e analisi", "Motore completo, DDS/BEN, replay, valore atteso e campo condiviso", "Robot, tornei, analisi e grandi campi competitivi", "Assente senza sviluppo specialistico"),
            ("Istruttori", "Classi, compiti, chat, lavagna, studio, generazione mani e tavolo live", "BBO offre tavoli didattici e stanze di pratica; Funbridge tornei/community", "Classi e contenuti sì; tavolo bridge no"),
            ("Dati e controllo", "Contenuti FIGB, export e replica del sorgente previsti dal contratto", "Servizio controllato dal fornitore", "Configurabilità dipendente dal piano e dal fornitore"),
            ("Personalizzazione", "Regole, flussi e contenuti adattabili alla Federazione italiana", "Prodotto standard globale", "Ampia sul corso, limitata sul motore di gioco"),
            ("Costo contrattuale", "Licenza e servizi gratuiti fino al 2029; licenza senza scadenza", "Freemium/abbonamenti e servizi del fornitore", "Canone di piattaforma più sviluppo bridge"),
            ("Scala comunitaria", "1.095 utenti registrati nella rilevazione", "BBO dichiara oltre 1 milione di visitatori mensili", "Dipende dalla comunità del cliente"),
        ],
        [1.30, 2.00, 1.80, 1.40],
        font_size=8.4,
    )
    add_table_note(doc, "Fonti competitor: pagine ufficiali Funbridge e Bridge Base Online consultate il 15 agosto 2026; collegamenti in appendice.")

    add_section_heading(doc, "4.1 Vantaggi specifici per la Federazione", 2)
    for text in [
        "Coerenza istituzionale: corsi, linguaggio, circoli, scuole e istruttori possono seguire un unico modello nazionale.",
        "Continuità del percorso: l'utente passa dalla teoria alla mano giocata, riceve un risultato, accumula progressione e può tornare sull'errore.",
        "Capacità di intervento: una funzione richiesta dalla Federazione può diventare parte del prodotto, anziché dipendere dalla roadmap di un SaaS globale.",
        "Controllo dei contenuti e dei dati: il contratto assegna a FIGB i propri materiali e ne garantisce l'esportazione; prevede inoltre una replica del codice.",
        "Costo di accesso: il pacchetto è gratuito nella fase di adozione e la licenza non ha una scadenza prefissata.",
    ]:
        add_bullet(doc, text, bullet_id)
    add_callout(doc, "DOVE I CONCORRENTI RESTANO PIÙ FORTI", "BBO e Funbridge hanno scala internazionale, rete di giocatori, multilingua, presenza Android e sistemi robot consolidati. BridgeLab deve puntare sulla specializzazione federale, non imitare la loro scala generalista.", fill=LIGHT_RED, accent=RED)

    add_section_heading(doc, "5. Maturità dell'infrastruttura — spiegata senza tecnicismi", 1)
    add_body(doc, "L'infrastruttura è l'insieme dei meccanismi che consentono alla piattaforma di funzionare, proteggere i dati, distribuire aggiornamenti e recuperare da un guasto. BridgeLab presenta una base produttiva credibile: non è un insieme di pagine isolate, ma un sistema a strati con controlli automatici e regole di accesso ai dati.")

    add_table(
        doc,
        ["Strato", "Che cosa fa", "Maturità osservata"],
        [
            ("Esperienza utente", "Web app installabile, navigazione, gioco, lezioni e portali", "81 pagine; PWA; build di produzione riuscita"),
            ("Logica BridgeLab", "Regole di gioco, progressione, quiz, sfide e strumenti scuola", "Moduli separati e testati; dipendenze DDS/BEN controllate"),
            ("Identità e dati", "Login, profili, corsi, risultati, classi e consensi", "47 tabelle; protezione per riga su 47/47; 91 policy"),
            ("Servizi specialistici", "Motore BEN, double-dummy, email e monitoraggio", "Timeout, guardie e fallback; dipendenza da servizi esterni"),
            ("Controllo qualità", "Verifica che gli aggiornamenti non rompano il prodotto", "990 test superati; 35 E2E definiti; 2 workflow CI"),
            ("Continuità", "Ricostruzione, passaggio a terzi e disponibilità del codice", "Baseline schema allineata; replica del sorgente prevista dal contratto"),
        ],
        [1.35, 2.65, 2.50],
        font_size=9.0,
    )

    add_section_heading(doc, "5.1 Valutazione sintetica", 2)
    add_table(
        doc,
        ["Area", "Punteggio / 5", "Lettura"],
        [
            ("Separazione delle responsabilità", "4", "Buona modularità; alcune pagine restano molto lunghe"),
            ("Modello dati", "4", "Integrità e protezione diffuse"),
            ("Autenticazione e autorizzazione", "4", "Controlli lato server e RLS"),
            ("Sicurezza", "3", "Impianto solido; segnalazioni automatiche da revisionare"),
            ("Gestione degli errori", "4", "Monitoraggio, timeout e fallback"),
            ("Prestazioni e caching", "3", "Buone tecniche; carico massimo non misurato"),
            ("Accessibilità", "4", "Controlli e test dedicati"),
            ("Internazionalizzazione", "1", "Prodotto sostanzialmente solo italiano"),
            ("Test", "4", "Suite ampia; coverage percentuale non determinabile"),
            ("CI/CD", "4", "Controlli principali e verifiche notturne"),
            ("Documentazione", "3", "Presente ma alcuni numeri sono disallineati"),
            ("Passaggio a un team terzo", "3", "Possibile, con rischio da conoscenza concentrata"),
        ],
        [2.55, 1.05, 2.90],
        numeric_cols=(1,),
        font_size=8.8,
    )
    add_callout(doc, "GIUDIZIO", "Media 3,42/5: infrastruttura adeguata a un servizio federale operativo e in crescita, ma non ancora dimostrata come infrastruttura critica con SLA, disaster recovery e carichi elevati certificati.", fill=LIGHT_GREEN, accent=GREEN)

    add_section_heading(doc, "5.2 Rischi e priorità", 2)
    for text in [
        "Concentrazione della conoscenza: 328 commit su 329 provengono dallo sviluppatore principale. La replica del codice riduce il rischio materiale, ma non trasferisce automaticamente il know-how.",
        "Capacità sotto carico: non è stato eseguito un load test; il massimo numero di utenti simultanei sostenibile è non determinabile.",
        "Continuità operativa: SLA, tempi di ripristino e prova di disaster recovery sono non determinabili e vanno formalizzati.",
        "Database: gli advisor segnalano ottimizzazioni su indici, policy e funzioni privilegiate. Sono rilievi tecnici da gestire, non vulnerabilità automaticamente dimostrate.",
        "Collaudo separato: alcuni test notturni puntano alla produzione; è preferibile un ambiente di staging dedicato.",
        "Canali: Android registra zero utenti nella dashboard; l'internazionalizzazione è assente.",
    ]:
        add_bullet(doc, text, bullet_id)

    add_section_heading(doc, "6. Valore della licenza gratuita", 1)
    add_body(doc, "La bozza contrattuale non trasferisce la proprietà intellettuale del nucleo BridgeLab. Concede però a FIGB un'esclusiva italiana per le finalità istituzionali e didattiche, senza scadenza prefissata e senza recesso unilaterale del licenziante salvo grave inadempimento. Hosting e manutenzione sono gratuiti fino al 31 dicembre 2029; il contratto disciplina inoltre contenuti FIGB, esportazione dei dati, replica del codice e diritti di continuità in eventi specifici.")
    add_callout(doc, "DISTINZIONE", "Il prezzo contrattuale è €0. Il valore economico è invece il beneficio che FIGB riceve senza pagarlo. Non coincide né con il costo storico di sviluppo né con il costo massimo di sostituzione.")

    add_section_heading(doc, "6.1 Metodo di valutazione", 2)
    add_body(doc, "Il metodo principale usa il costo evitato di una piattaforma comparabile per l'apprendimento del bridge. Per i primi tre anni FIGB non sostiene né canone né costi di hosting/manutenzione. Dopo il periodo gratuito si sottrae l'onere operativo necessario a mantenere una soluzione propria. Poiché il diritto è giuridicamente senza termine ma il software può diventare obsoleto, la stima centrale usa una vita economica di 10 anni, non una rendita infinita.")
    add_table(
        doc,
        ["Assunzione", "−30%", "Base", "+30%"],
        [
            ("Utenti attivi mensili sostenibili", "197,4", "282", "366,6"),
            ("Costo operativo annuo dopo il periodo gratuito", "€15.452,16", "€22.074,52", "€28.696,87"),
            ("Vita economica", "7 anni", "10 anni", "13 anni"),
            ("Tasso di attualizzazione", "3,5%", "5,0%", "6,5%"),
        ],
        [2.85, 1.15, 1.20, 1.30],
        numeric_cols=(1, 2, 3),
        font_size=9.1,
    )
    add_table_note(doc, "Prezzo osservato, non assunto: Funbridge Premium €149,99 per utente/anno. Orizzonte all-inclusive: 3 anni, coerente con la perizia e ricompreso nel termine contrattuale del 31 dicembre 2029.")

    add_section_heading(doc, "6.2 Risultato", 2)
    add_table(
        doc,
        ["Componente, scenario centrale", "Valore attuale"],
        [
            ("Primi tre anni: licenza + hosting + manutenzione", "€115.185,71"),
            ("Licenza residua dal 4º al 10º anno, al netto dei costi operativi", "€101.082,72"),
            ("Valore complessivo centrale", "€216.268,43"),
        ],
        [4.65, 1.85],
        numeric_cols=(1,),
    )
    add_picture_with_alt(
        doc,
        value_chart,
        width=Inches(6.25),
        alt_text="Grafico a colonne del valore attuale del contratto negli scenari basso, centrale e alto, distinto tra primi tre anni all-inclusive e licenza residua.",
    )
    add_caption(doc, "Figura 3 — Valore attuale: scenario basso €81.000, centrale €216.268, alto €450.602.")

    add_table(
        doc,
        ["Misura", "Valore", "Uso corretto"],
        [
            ("Beneficio 3 anni non attualizzato", "€88.824–€164.959; base €126.891,54", "Valore iniziale all-inclusive"),
            ("Valore totale prudenziale", "€81.000–€450.602; base €216.268,43", "Range consigliato per il Consiglio"),
            ("Perpetuità teorica centrale", "€464.567,65", "Solo controllo; non raccomandata come cifra ufficiale"),
            ("Costo di riproduzione a effort base", "€568.400–€893.200", "Tetto tecnico, non valore d'uso"),
            ("Costo fornitore a perimetro base", "€974.400–€1.827.000", "Controfattuale di gara, non valore conferito"),
        ],
        [2.15, 2.20, 2.15],
        font_size=8.9,
    )
    add_callout(doc, "VALORE DA COMUNICARE", "Il contratto attribuisce gratuitamente a FIGB un beneficio economico stimato in circa €216.000, con un intervallo prudenziale di €81.000–€451.000. Il valore comprende i primi tre anni all-inclusive e il valore residuo della licenza italiana a tempo indeterminato.", fill=LIGHT_GREEN, accent=GREEN)
    add_body(doc, "La perpetuità giuridica non equivale a vita economica infinita. Inoltre l'esclusiva è limitata alle finalità istituzionali e didattiche e FIGB non può commercializzarla o sublicenziarla liberamente. In caso di recesso volontario di FIGB è previsto un periodo transitorio; i diritti perpetui sul codice si attivano pienamente negli eventi descritti dall'articolo 5. Questi limiti sono già riflessi nell'approccio prudente.")

    add_section_heading(doc, "7. Conclusioni per il Consiglio", 1)
    add_body(doc, "BridgeLab combina tre asset difficili da separare: piattaforma tecnica, patrimonio didattico e processo federale. La sua utilità è dimostrata da 1.095 utenti, oltre 8.221 ore di utilizzo e 60.773 partite; la sua profondità è dimostrata da un catalogo articolato e da strumenti che coprono allievo, istruttore, classe, gioco e amministrazione.")
    add_body(doc, "L'infrastruttura è matura per l'operatività attuale, con protezione dei dati, test, build e automazioni. Per trasformarla in un servizio istituzionale pienamente presidiato occorre però ridurre la dipendenza da una sola persona, formalizzare continuità e livelli di servizio, introdurre staging e load test e mantenere allineata la documentazione.")
    add_body(doc, "Il valore economico più difendibile del contratto non è il milione potenzialmente necessario per ricommissionare tutto a un fornitore. È il beneficio d'uso della licenza e dei servizi, attualizzato e corretto per i costi futuri: €216.268 nello scenario centrale, con range prudenziale €81.000–€450.602.")

    add_section_heading(doc, "Appendice A — Metodo e fonti", 1)
    add_body(doc, "Ogni dato numerico della relazione deriva dai comandi o dai file indicati di seguito. Le statistiche di utilizzo sono riportate esclusivamente in forma aggregata.")

    add_section_heading(doc, "A.1 Repository ed effort", 2)
    add_body(doc, "Comando eseguito:", bold_lead="Comando eseguito:", keep=True)
    add_callout(doc, "COMANDO", "node audit-bridgelab/collect-git-metrics.mjs", fill=LIGHT_GRAY, accent=MUTED)
    add_body(doc, "Output utilizzato: 329 commit; autore principale 328; 57 giorni attivi; 102 sessioni; 127,84 ore; fasce 00–05: 58, 06–11: 66, 12–17: 102, 18–23: 103; giorni lun–dom: 32, 57, 60, 56, 38, 36, 50; weekend 86; fuori finestra 09–18: 224.")

    add_section_heading(doc, "A.2 Inventario di prodotto", 2)
    add_callout(doc, "COMANDO", "npx tsx audit-bridgelab/collect-static-product-metrics.ts", fill=LIGHT_GRAY, accent=MUTED)
    add_body(doc, "Output utilizzato: 4 corsi, 49 lezioni, 168 moduli nel sorgente, 943 blocchi, 5 tipologie quiz, 111 domande di comprensione, 340 definizioni di mani, 13 badge, 10 obiettivi segreti, 36 livelli, 81 pagine, 102 componenti e 13 handler API. I valori vivi di 199 moduli e 1.504 mani generate derivano dalle query aggregate documentate nella perizia tecnica.")

    add_section_heading(doc, "A.3 Dashboard di utilizzo", 2)
    add_body(doc, "File letto: /Users/albertogiovannigerli/.codex/attachments/5288cdbf-60a7-4ba7-9dc2-9161305ac1b3/pasted-text.txt. Sono state usate soltanto le sezioni aggregate iniziali della dashboard; l'elenco individuale degli utenti non è stato riportato né elaborato nella relazione.")
    add_body(doc, "Output aggregato utilizzato: 1.095 utenti; 3 nuovi oggi, 12 in 7 giorni, 49 in 30 giorni; 79 attivi oggi, 133 in 7 giorni; 8.221 h 48 m complessive; 450 minuti medi; retention 7 giorni 11%; 60.773 partite; 665 giocatori; 16 istruttori; 16 classi; 52 allievi.")

    add_section_heading(doc, "A.4 Valore contrattuale", 2)
    add_callout(doc, "COMANDO", "node audit-bridgelab/contract-value-calculations.mjs", fill=LIGHT_GRAY, accent=MUTED)
    add_body(doc, "Output centrale: utenti attivi 282; costo operativo annuo €22.074,52; vita economica 10 anni; tasso 5%; primi tre anni €115.185,71; licenza residua €101.082,72; totale €216.268,43. Output di sensibilità: €81.000,21–€450.601,59. Perpetuità teorica centrale: €464.567,65.")

    add_section_heading(doc, "A.5 Qualità e infrastruttura", 2)
    add_body(doc, "Comandi eseguiti nella perizia tecnica: npm test; npx tsc --noEmit; npx eslint src --max-warnings 0; npm run build; npx playwright test --list; python3 deploy/ben-railway/test_guard.py; npm run schema:check. Risultati utilizzati: 990 test passati, 5 saltati; TypeScript e lint con esito positivo; 85 pagine generate in build; 35 E2E elencati; 11 controlli BEN passati; schema allineato su 4.514 righe.")

    add_section_heading(doc, "A.6 Fonti pubbliche sui concorrenti", 2)
    sources = [
        ("Funbridge — apprendimento e modulo introduttivo", "https://funbridge.com/learn-bridge"),
        ("Funbridge — gioco, IA, analisi e modalità", "https://funbridge.com/"),
        ("Bridge Base Online — caratteristiche", "https://www.bridgebase.com/intro/features.php"),
        ("Bridge Base Online — pratica e tavoli didattici", "https://news.bridgebase.com/about-practice/"),
        ("Bridge Base Online — dimensione della comunità", "https://news.bridgebase.com/about/"),
        ("Funbridge — offerte e prezzo Premium", "https://funbridge.com/fr/offres"),
    ]
    for label, url in sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        add_hyperlink(p, label, url)
        set_run_font(p.add_run(" — consultato il 15 agosto 2026"), size=9.5, color=MUTED)

    add_section_heading(doc, "A.7 Limiti", 2)
    add_body(doc, "La relazione non è un audit legale, contabile, di sicurezza o di accessibilità. Il contratto esaminato è una bozza con revisioni e campi non completati; il valore presuppone la firma del testo sostanzialmente invariato. La durata giuridica della licenza non garantisce una durata economica infinita. La retention deriva dalla definizione della dashboard e non è stata ricalcolata sui dati individuali. Le ore Git sono un limite inferiore. Il carico massimo, gli SLA, i tempi di ripristino, la configurazione effettiva dei backup e il valore contabile della licenza sono non determinabili dalle fonti esaminate.")

    doc.core_properties.title = "Relazione BridgeLab per il Consiglio Federale"
    doc.core_properties.subject = "Funzionalità, adozione, effort, infrastruttura e valore della licenza"
    doc.core_properties.author = "Analisi tecnica indipendente"
    doc.core_properties.keywords = "BridgeLab, FIGB, piattaforma, licenza, infrastruttura, utenti"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
