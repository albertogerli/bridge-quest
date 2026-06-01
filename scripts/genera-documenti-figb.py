#!/usr/bin/env python3
"""
Genera i documenti FIGB per la formalizzazione Bridge LAB:
1. Lettera al Segretario Generale
2. Accordo art. 28 GDPR (Titolare/Responsabile)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_style(doc):
    """Imposta lo stile base del documento."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x3D, 0xA5)
    return h


def genera_lettera_segretario():
    doc = Document()
    set_style(doc)

    # Margini
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Intestazione
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Alberto Giovanni Gerli\nVice Presidente FIGB\nTourbillon Tech S.r.l.")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Spett.le Segretario Generale\nFederazione Italiana Gioco Bridge")
    run.font.size = Pt(11)
    run.bold = True

    doc.add_paragraph()

    # Oggetto
    p = doc.add_paragraph()
    run = p.add_run("Oggetto: ")
    run.bold = True
    run = p.add_run("Piattaforma Bridge LAB (bridgelab.it) — Formalizzazione rapporto FIGB / Tourbillon Tech S.r.l.")

    doc.add_paragraph()

    # Corpo
    doc.add_paragraph(
        "Caro Segretario,"
    )

    doc.add_paragraph(
        "come noto, ho sviluppato e messo online la piattaforma didattica Bridge LAB "
        "(bridgelab.it), che ospita i 4 corsi della Commissione Insegnamento (Fiori, Quadri, "
        "Cuori — Gioco della Carta, Cuori — Dichiarazione) in formato interattivo e gamificato, "
        "per un totale di 49 lezioni e oltre 168 moduli didattici."
    )

    doc.add_paragraph(
        "Lo sviluppo software, la gestione tecnica e l'hosting della piattaforma sono "
        "interamente a carico mio e della mia società Tourbillon Tech S.r.l., "
        "a titolo completamente gratuito, come contributo personale alla Federazione."
    )

    doc.add_paragraph(
        "Per regolarizzare la situazione sotto il profilo formale e di conformità al "
        "Regolamento UE 2016/679 (GDPR), chiedo che il Consiglio Federale voglia "
        "deliberare quanto segue."
    )

    # Punti delibera
    add_heading_styled(doc, "Punti da sottoporre al Consiglio Federale", level=2)

    items = [
        (
            "Presa d'atto e approvazione",
            "della piattaforma Bridge LAB (bridgelab.it) come strumento didattico "
            "ufficiale della FIGB, sviluppato nell'ambito della Commissione Insegnamento."
        ),
        (
            "Presa d'atto della gratuità",
            "che lo sviluppo software, la gestione tecnica e l'hosting della piattaforma "
            "sono forniti a titolo gratuito da Tourbillon Tech S.r.l. (legale rappresentante: "
            "Alberto Giovanni Gerli, Vice Presidente FIGB), senza alcun corrispettivo "
            "né diretto né indiretto a carico della FIGB."
        ),
        (
            "Nomina a Responsabile del trattamento",
            "di Tourbillon Tech S.r.l. quale Responsabile del trattamento dei dati personali "
            "degli utenti della piattaforma, ai sensi dell'art. 28 del Regolamento UE 2016/679 "
            "(GDPR), con autorizzazione alla sottoscrizione del relativo accordo tra FIGB "
            "(Titolare) e Tourbillon Tech S.r.l. (Responsabile). La bozza di accordo è "
            "allegata alla presente."
        ),
        (
            "Titolarità dei contenuti didattici",
            "conferma che i contenuti didattici (testi dei corsi, materiale delle lezioni, "
            "video) restano di proprietà della FIGB. Il codice sorgente della piattaforma "
            "resta nella disponibilità di Tourbillon Tech S.r.l."
        ),
        (
            "Utilizzo del marchio FIGB",
            "autorizzazione all'utilizzo del nome, del logo e del marchio FIGB sulla "
            "piattaforma Bridge LAB, esclusivamente per le finalità didattiche sopra descritte."
        ),
    ]

    for i, (title, desc) in enumerate(items, 1):
        p = doc.add_paragraph()
        p.style = doc.styles["List Number"]
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # Conflitto di interesse
    add_heading_styled(doc, "Dichiarazione sul conflitto d'interesse", level=2)

    doc.add_paragraph(
        "Essendo io Vice Presidente della FIGB e al contempo legale rappresentante "
        "di Tourbillon Tech S.r.l., dichiaro che:"
    )

    conflict_items = [
        "il servizio è fornito a titolo integralmente gratuito, senza alcun corrispettivo né diretto né indiretto;",
        "non sussiste alcun rapporto economico tra FIGB e Tourbillon Tech S.r.l. relativamente alla piattaforma Bridge LAB;",
        "mi asterrò dalla votazione su questa delibera in sede di Consiglio Federale.",
    ]

    for item in conflict_items:
        p = doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()

    doc.add_paragraph(
        "Resto a disposizione per qualsiasi chiarimento e per predisporre la "
        "documentazione necessaria."
    )

    doc.add_paragraph()

    p = doc.add_paragraph("Con i migliori saluti,")
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Alberto Giovanni Gerli")
    run.bold = True
    p2 = doc.add_paragraph("Vice Presidente FIGB")
    p3 = doc.add_paragraph("Tourbillon Tech S.r.l.")

    # Salva
    path = os.path.join(OUTPUT_DIR, "Lettera-Segretario-FIGB-BridgeLAB.docx")
    doc.save(path)
    print(f"✅ Lettera salvata: {path}")
    return path


def genera_accordo_art28():
    doc = Document()
    set_style(doc)

    # Margini
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Titolo
    title = doc.add_heading("ACCORDO PER IL TRATTAMENTO DEI DATI PERSONALI", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x00, 0x3D, 0xA5)

    subtitle = doc.add_paragraph("ai sensi dell'art. 28 del Regolamento UE 2016/679 (GDPR)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Parti
    add_heading_styled(doc, "TRA", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Federazione Italiana Gioco Bridge (FIGB)")
    run.bold = True
    p.add_run(
        ", con sede in Via Tuscolana 65, 00182 Roma, C.F. 80208250585, "
        "in persona del Presidente pro tempore, di seguito denominata "
    )
    run = p.add_run('"Titolare"')
    run.bold = True

    doc.add_paragraph("E")

    p = doc.add_paragraph()
    run = p.add_run("Tourbillon Tech S.r.l.")
    run.bold = True
    p.add_run(
        ", in persona del legale rappresentante Alberto Giovanni Gerli, "
        "di seguito denominata "
    )
    run = p.add_run('"Responsabile"')
    run.bold = True

    doc.add_paragraph()

    # Premesse
    add_heading_styled(doc, "PREMESSO CHE", level=2)

    premesse = [
        "il Titolare è titolare del trattamento dei dati personali degli utenti della "
        "piattaforma didattica \"Bridge LAB\" (bridgelab.it);",
        "il Responsabile ha sviluppato la piattaforma Bridge LAB e ne cura la gestione "
        "tecnica e l'hosting a titolo completamente gratuito;",
        "il Responsabile, nell'ambito dell'erogazione dei servizi tecnici, tratta dati "
        "personali per conto del Titolare;",
        "le parti intendono regolare i rispettivi obblighi in conformità all'art. 28 "
        "del Regolamento UE 2016/679 (GDPR).",
    ]

    for item in premesse:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("SI CONVIENE E SI STIPULA QUANTO SEGUE")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Articoli
    articoli = [
        (
            "Oggetto",
            [
                "Il presente accordo disciplina il trattamento dei dati personali che il "
                "Responsabile effettua per conto del Titolare nell'ambito della gestione "
                "tecnica della piattaforma Bridge LAB.",
                "Il Responsabile presta i propri servizi a titolo completamente gratuito."
            ]
        ),
        (
            "Tipologia di dati e interessati",
            [
                "I dati personali oggetto del trattamento comprendono: dati di registrazione "
                "(email, nome visualizzato, fascia d'età, ASD di appartenenza), dati di utilizzo "
                "(progressi, punteggi, risultati di gioco), dati tecnici (indirizzo IP, browser, "
                "sistema operativo).",
                "Gli interessati sono gli utenti registrati e i visitatori della piattaforma Bridge LAB."
            ]
        ),
        (
            "Obblighi del Responsabile",
            [
                "Il Responsabile si impegna a:",
                "a) trattare i dati personali soltanto su istruzione documentata del Titolare;",
                "b) garantire che le persone autorizzate al trattamento si siano impegnate alla riservatezza;",
                "c) adottare tutte le misure di sicurezza tecniche e organizzative adeguate ai sensi dell'art. 32 GDPR;",
                "d) assistere il Titolare nel garantire il rispetto degli obblighi di cui agli artt. 32-36 GDPR;",
                "e) su scelta del Titolare, cancellare o restituire tutti i dati personali al termine della prestazione;",
                "f) mettere a disposizione del Titolare tutte le informazioni necessarie per dimostrare il rispetto del presente accordo."
            ]
        ),
        (
            "Sub-responsabili",
            [
                "Il Titolare autorizza il Responsabile ad avvalersi dei seguenti sub-responsabili:",
                "• Supabase Inc. — database e autenticazione (server UE, Francoforte);",
                "• Vercel Inc. — hosting applicazione e analytics;",
                "• Google LLC / YouTube — video didattici embedded.",
                "Il Responsabile informerà tempestivamente il Titolare in caso di modifiche relative "
                "all'aggiunta o alla sostituzione di sub-responsabili."
            ]
        ),
        (
            "Misure di sicurezza",
            [
                "Il Responsabile adotta le seguenti misure di sicurezza:",
                "a) crittografia dei dati in transito (HTTPS/TLS) e a riposo;",
                "b) autenticazione degli utenti tramite protocolli sicuri (Supabase Auth);",
                "c) accesso ai dati limitato al solo personale autorizzato del Responsabile;",
                "d) backup periodici dei dati;",
                "e) Row Level Security (RLS) sul database per l'isolamento dei dati tra utenti."
            ]
        ),
        (
            "Violazione dei dati (Data Breach)",
            [
                "Il Responsabile notificherà al Titolare qualsiasi violazione dei dati personali "
                "senza ingiustificato ritardo e comunque entro 48 ore dal momento in cui ne è "
                "venuto a conoscenza, fornendo tutte le informazioni necessarie ai sensi "
                "dell'art. 33 GDPR."
            ]
        ),
        (
            "Durata",
            [
                "Il presente accordo ha durata pari alla prestazione dei servizi tecnici da parte "
                "del Responsabile e potrà essere risolto da ciascuna delle parti con preavviso "
                "scritto di 90 giorni.",
                "In caso di cessazione, il Responsabile restituirà o cancellerà tutti i dati personali "
                "entro 30 giorni, salvo obblighi di legge."
            ]
        ),
        (
            "Gratuità e assenza di conflitto",
            [
                "Le parti danno atto che i servizi del Responsabile sono prestati a titolo "
                "integralmente gratuito, quale contributo personale del legale rappresentante "
                "del Responsabile (nella sua qualità di Vice Presidente della FIGB) alla "
                "Federazione.",
                "Non sussiste alcun rapporto economico tra le parti relativamente "
                "alla piattaforma Bridge LAB."
            ]
        ),
        (
            "Legge applicabile e foro competente",
            [
                "Il presente accordo è regolato dalla legge italiana. Per qualsiasi controversia "
                "sarà competente il Foro di Roma."
            ]
        ),
    ]

    for i, (title, paragraphs) in enumerate(articoli, 1):
        add_heading_styled(doc, f"Art. {i} — {title}", level=2)
        for text in paragraphs:
            doc.add_paragraph(text)

    doc.add_paragraph()

    # Firme
    add_heading_styled(doc, "FIRME", level=2)

    p = doc.add_paragraph()
    p.add_run("Luogo e data: _______________________")

    doc.add_paragraph()

    # Tabella firme
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cells = table.rows[0].cells
    r = cells[0].paragraphs[0].add_run("Per il Titolare")
    r.bold = True
    r.font.size = Pt(10)
    r = cells[1].paragraphs[0].add_run("Per il Responsabile")
    r.bold = True
    r.font.size = Pt(10)

    cells = table.rows[1].cells
    r = cells[0].paragraphs[0].add_run("Federazione Italiana Gioco Bridge")
    r.font.size = Pt(10)
    r = cells[1].paragraphs[0].add_run("Tourbillon Tech S.r.l.")
    r.font.size = Pt(10)

    cells = table.rows[2].cells
    r = cells[0].paragraphs[0].add_run("Il Presidente")
    r.font.size = Pt(10)
    r = cells[1].paragraphs[0].add_run("Il Legale Rappresentante")
    r.font.size = Pt(10)

    cells = table.rows[3].cells
    cells[0].paragraphs[0].add_run("\n\n_______________________").font.size = Pt(10)
    r = cells[1].paragraphs[0].add_run("\n\nAlberto Giovanni Gerli\n\n_______________________")
    r.font.size = Pt(10)

    # Remove borders from table
    from docx.oxml.ns import qn
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn("w:tcBorders"))
            if tcBorders is not None:
                tcPr.remove(tcBorders)

    # Salva
    path = os.path.join(OUTPUT_DIR, "Accordo-Art28-GDPR-FIGB-TourbillonTech.docx")
    doc.save(path)
    print(f"✅ Accordo art. 28 salvato: {path}")
    return path


if __name__ == "__main__":
    print("📄 Generazione documenti FIGB / Tourbillon Tech...")
    print()
    genera_lettera_segretario()
    genera_accordo_art28()
    print()
    print("✅ Tutti i documenti generati nella cartella docs/")
