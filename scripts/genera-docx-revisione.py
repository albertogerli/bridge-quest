#!/usr/bin/env python3
"""
Genera il documento Word con TUTTI i contenuti didattici di BridgeLab, per la
revisione di un esperto federale.

PERCHÉ NON RIUSA scripts/generate-review-doc.ts
Quello script legge da `src/data/*.ts`, cioè dal seed. I contenuti in
produzione sono su Supabase e DIVERGONO dal seed (vedi CLAUDE.md): un esperto
che revisionasse il seed correggerebbe testi che nessun allievo vede. Qui si
parte dal dump del database.

COSA CONTIENE
  * tutti i moduli di tutte le lezioni, blocco per blocco, nell'ordine in cui
    l'allievo li incontra;
  * per ogni quiz: le opzioni, quale sia data per corretta e la spiegazione —
    senza questi tre elementi insieme non è verificabile;
  * i diagrammi di tutte le smazzate nella disposizione Nord/Ovest/Est/Sud,
    con i punti onori calcolati da noi: se un testo dichiara punti diversi,
    lo scarto si vede a occhio;
  * glossario ed eserciziario.

Ogni voce riporta lezione e modulo di provenienza, così una correzione può
essere riportata al punto giusto del database senza cercarla.

USO
    node  ...dump...            # produce _dump-contenuti.json
    python3 scripts/genera-docx-revisione.py
"""

import json
import os
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DUMP = os.path.join(ROOT, "_dump-contenuti.json")
OUT = os.path.join(ROOT, "Revisione_Contenuti_BridgeLab.docx")

BLU = RGBColor(0x00, 0x3D, 0xA5)
ROSSO = RGBColor(0xB9, 0x1C, 0x1C)
GRIGIO = RGBColor(0x66, 0x66, 0x66)
VERDE = RGBColor(0x0B, 0x6B, 0x3A)

SEMI = [("spade", "♠"), ("heart", "♥"), ("diamond", "♦"), ("club", "♣")]
ORDINE = ["A", "K", "Q", "J", "10", "9", "8", "7", "6", "5", "4", "3", "2"]
PUNTI = {"A": 4, "K": 3, "Q": 2, "J": 1}
POS_IT = {"north": "Nord", "east": "Est", "south": "Sud", "west": "Ovest"}
VULN_IT = {"none": "Nessuna", "ns": "N-S", "ew": "E-O", "both": "Tutti"}

# Etichette dei tipi di blocco, così il revisore sa che cosa sta guardando.
TIPI = {
    "text": "Testo",
    "heading": "Titolo",
    "rule": "Regola",
    "tip": "Consiglio",
    "example": "Esempio",
    "quiz": "Quiz",
    "true-false": "Vero/Falso",
    "bid-select": "Scelta della dichiarazione",
    "card-select": "Scelta della carta",
    "hand-eval": "Valutazione della mano",
}


# ─── utilità ────────────────────────────────────────────────────────────────

def carte_per_seme(mano):
    """Carte divise per seme, in ordine decrescente."""
    out = {}
    for chiave, _ in SEMI:
        carte = [c["rank"] for c in mano if c["suit"] == chiave]
        carte.sort(key=lambda r: ORDINE.index(r) if r in ORDINE else 99)
        out[chiave] = carte
    return out


def punti_onori(mano):
    return sum(PUNTI.get(c["rank"], 0) for c in mano)


def par(doc, testo, *, size=11, bold=False, italic=False, color=None,
        space_after=4, align=None, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(testo)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    return p


def etichetta(doc, testo, colore):
    """Riga breve che introduce un blocco (es. «QUIZ · lezione 3, modulo 3-1»)."""
    return par(doc, testo, size=8, bold=True, color=colore, space_after=2)


# ─── diagramma di una smazzata ──────────────────────────────────────────────

def diagramma(doc, smazzata):
    """
    Disegna la mano nella disposizione del tavolo: Nord in alto, Ovest ed Est
    ai lati, Sud in basso. È la forma in cui ogni giocatore di bridge legge una
    mano; un elenco lineare costringerebbe il revisore a ricostruirla a mente.
    """
    mani = smazzata.get("hands") or {}
    tab = doc.add_table(rows=3, cols=3)
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER

    def scrivi(cella, posizione):
        mano = mani.get(posizione) or []
        cella.text = ""
        intest = cella.paragraphs[0]
        r = intest.add_run(f"{POS_IT[posizione]} ({punti_onori(mano)} PO)")
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = GRIGIO
        intest.paragraph_format.space_after = Pt(1)
        per_seme = carte_per_seme(mano)
        for chiave, simbolo in SEMI:
            p = cella.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(f"{simbolo} {' '.join(per_seme[chiave]) or '—'}")
            run.font.size = Pt(9)
            run.font.name = "Consolas"
            # Rossi e neri come al tavolo: aiuta a cogliere subito la sagoma.
            run.font.color.rgb = ROSSO if chiave in ("heart", "diamond") else RGBColor(0x1A, 0x1A, 0x2E)

    scrivi(tab.cell(0, 1), "north")
    scrivi(tab.cell(1, 0), "west")
    scrivi(tab.cell(1, 2), "east")
    scrivi(tab.cell(2, 1), "south")

    centro = tab.cell(1, 1)
    centro.text = ""
    lead = smazzata.get("opening_lead") or {}
    simbolo_lead = dict(SEMI).get(lead.get("suit"), "")
    righe = [
        smazzata.get("contract") or "—",
        f"Dich. {POS_IT.get(smazzata.get('declarer'), '—')}",
        f"Vuln. {VULN_IT.get(smazzata.get('vulnerability'), '—')}",
        f"Attacco {simbolo_lead}{lead.get('rank', '')}" if lead else "",
    ]
    for i, riga in enumerate([r for r in righe if r]):
        p = centro.paragraphs[0] if i == 0 else centro.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(riga)
        run.font.size = Pt(8)
        run.bold = i == 0
        run.font.color.rgb = BLU if i == 0 else GRIGIO

    # Somma di controllo: se non fa 40 o 52, la mano è corrotta.
    totale_po = sum(punti_onori(mani.get(p) or []) for p in POS_IT)
    totale_carte = sum(len(mani.get(p) or []) for p in POS_IT)
    if totale_po != 40 or totale_carte != 52:
        par(doc, f"⚠ Anomalia: {totale_carte} carte e {totale_po} punti in totale "
                 f"(attesi 52 e 40).", size=9, bold=True, color=ROSSO)
    doc.add_paragraph()


# ─── blocchi di contenuto ───────────────────────────────────────────────────

def scrivi_blocco(doc, blocco, rif):
    tipo = blocco.get("type", "?")
    nome = TIPI.get(tipo, tipo)
    contenuto = (blocco.get("content") or "").strip()

    if tipo == "heading":
        par(doc, contenuto, size=13, bold=True, color=BLU, space_after=3)
        return

    if tipo in ("text", "rule", "tip"):
        etichetta(doc, f"{nome.upper()} · {rif}", GRIGIO)
        par(doc, contenuto, size=10.5, space_after=8)
        return

    if tipo == "example":
        etichetta(doc, f"ESEMPIO · {rif}", GRIGIO)
        carte = (blocco.get("cards") or "").strip()
        if carte:
            par(doc, carte, size=10, bold=True, color=BLU, space_after=2)
        par(doc, contenuto, size=10.5, space_after=8)
        return

    # Quiz e varianti: opzioni, risposta data per corretta, spiegazione.
    etichetta(doc, f"{nome.upper()} · {rif}", BLU)
    par(doc, contenuto, size=10.5, bold=True, space_after=3)

    opzioni = blocco.get("options") or []
    corretta = blocco.get("correctAnswer")
    if opzioni:
        for i, opzione in enumerate(opzioni):
            giusta = i == corretta
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(("✔ " if giusta else "○ ") + str(opzione))
            run.font.size = Pt(10)
            run.bold = giusta
            if giusta:
                run.font.color.rgb = VERDE
    elif tipo == "true-false":
        # Senza opzioni esplicite, correctAnswer 1 = Vero, 0 = Falso.
        risposta = "VERO" if corretta == 1 else "FALSO"
        par(doc, f"✔ Risposta data per corretta: {risposta}", size=10,
            bold=True, color=VERDE, space_after=2)

    spiegazione = (blocco.get("explanation") or "").strip()
    if spiegazione:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run("Spiegazione: " + spiegazione)
        run.font.size = Pt(9.5)
        run.italic = True
        run.font.color.rgb = GRIGIO
    else:
        doc.add_paragraph()


# ─── documento ──────────────────────────────────────────────────────────────

def main():
    with open(DUMP, encoding="utf-8") as f:
        dati = json.load(f)

    corsi = sorted(dati["courses"], key=lambda c: c.get("position") or 0)
    # `lessons.world_id` punta a un MONDO, non al corso: il collegamento passa
    # da `course_worlds`. Sbagliarlo non dà errore, dà un documento vuoto —
    # ed è esattamente quello che è successo alla prima stesura.
    mondo_a_corso = {w["id"]: w.get("course_id") for w in dati.get("worlds", [])}
    lezioni = dati["lessons"]
    moduli = dati["modules"]
    smazzate = dati["smazzate"]
    glossario = dati["glossary"]
    eserciziario = dati.get("eserc") or []

    per_lezione = {}
    for m in moduli:
        per_lezione.setdefault(m["lesson_id"], []).append(m)
    for v in per_lezione.values():
        v.sort(key=lambda m: m.get("position") or 0)

    # `smazzate.lesson_id` è una STRINGA mentre `lessons.id` è un intero:
    # raggrupparli senza normalizzare fa sparire tutti i diagrammi.
    def come_intero(v):
        try:
            return int(v)
        except (TypeError, ValueError):
            return None

    smazzate_per_lezione = {}
    for s in smazzate:
        smazzate_per_lezione.setdefault(come_intero(s.get("lesson_id")), []).append(s)
    for v in smazzate_per_lezione.values():
        v.sort(key=lambda s: s.get("board") or 0)

    doc = Document()
    for stile in ("Normal",):
        doc.styles[stile].font.name = "Calibri"
        doc.styles[stile].font.size = Pt(11)

    # Copertina
    par(doc, "BridgeLab", size=30, bold=True, color=BLU,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    par(doc, "Contenuti didattici — documento per la revisione", size=15,
        color=GRIGIO, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    par(doc, f"Estratto dai contenuti in produzione il "
             f"{date.today().strftime('%d/%m/%Y')}", size=10,
        color=GRIGIO, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    par(doc, "Che cosa contiene", size=13, bold=True, color=BLU)
    par(doc,
        f"{len(corsi)} corsi · {len(lezioni)} lezioni · {len(moduli)} moduli · "
        f"{len(smazzate)} smazzate · {len(glossario)} voci di glossario · "
        f"{len(eserciziario)} esercizi.", size=10.5, space_after=10)
    par(doc,
        "I testi provengono dal database di produzione, non dai file sorgente: "
        "sono esattamente quelli che gli allievi leggono oggi. Ogni voce riporta "
        "fra parentesi la lezione e il modulo di provenienza, così una correzione "
        "può essere riportata al punto esatto.", size=10.5, space_after=10)
    par(doc,
        "Per ogni quiz sono riportate tutte le opzioni, con un segno di spunta "
        "sulla risposta data per corretta, e la spiegazione mostrata all'allievo. "
        "Vanno verificate insieme: una spiegazione giusta accanto a una risposta "
        "sbagliata è l'errore più difficile da notare in produzione.",
        size=10.5, space_after=10)
    par(doc,
        "Nei diagrammi i punti onori accanto a ogni posizione sono ricalcolati da "
        "noi a partire dalle carte, non copiati dal testo: se un enunciato dichiara "
        "un valore diverso, la differenza è visibile subito.", size=10.5, space_after=10)
    par(doc,
        "Come annotare: usare i commenti di Word (Revisione → Nuovo commento) "
        "oppure le revisioni con traccia. Non serve riscrivere il testo corretto "
        "per intero: basta indicare che cosa non va.", size=10.5, space_after=10)

    anomalie = []

    # Corsi → lezioni → moduli
    for corso in corsi:
        doc.add_page_break()
        par(doc, corso.get("name") or "Corso", size=22, bold=True, color=BLU, space_after=2)
        if corso.get("subtitle"):
            par(doc, corso["subtitle"], size=11, color=GRIGIO, space_after=12)

        lezioni_corso = sorted(
            [l for l in lezioni
             if mondo_a_corso.get(l.get("world_id")) == corso.get("id")],
            key=lambda l: l.get("id") or 0,
        )
        for lezione in lezioni_corso:
            par(doc, f"Lezione {lezione['id']} — {lezione.get('title') or ''}",
                size=16, bold=True, space_after=2)
            if lezione.get("subtitle"):
                par(doc, lezione["subtitle"], size=10, italic=True, color=GRIGIO, space_after=8)

            for modulo in per_lezione.get(lezione["id"], []):
                par(doc, f"{modulo.get('title') or modulo['module_id']}",
                    size=13, bold=True, color=BLU, space_after=4)
                rif = f"lez. {lezione['id']}, mod. {modulo['module_id']}"
                blocchi = modulo.get("content")
                if isinstance(blocchi, list):
                    for blocco in blocchi:
                        if isinstance(blocco, dict):
                            scrivi_blocco(doc, blocco, rif)

            for s in smazzate_per_lezione.get(lezione["id"], []):
                par(doc, f"Smazzata {s.get('board')} — {s.get('title') or ''}",
                    size=12, bold=True, space_after=3)
                diagramma(doc, s)
                if s.get("commentary"):
                    par(doc, s["commentary"], size=10, space_after=10)
                mani = s.get("hands") or {}
                if sum(len(mani.get(p) or []) for p in POS_IT) != 52:
                    anomalie.append(f"smazzata {s.get('id')} (lezione {s.get('lesson_id')})")

    # Smazzate non collegate a una lezione
    lezioni_note = {l["id"] for l in lezioni}
    orfane = [s for s in smazzate if come_intero(s.get("lesson_id")) not in lezioni_note]
    if orfane:
        doc.add_page_break()
        par(doc, "Smazzate non collegate a una lezione", size=22, bold=True, color=BLU, space_after=10)
        for s in sorted(orfane, key=lambda x: (come_intero(x.get("lesson_id")) or 0, x.get("board") or 0)):
            par(doc, f"{s.get('id')} — {s.get('title') or ''}", size=12, bold=True, space_after=3)
            diagramma(doc, s)
            if s.get("commentary"):
                par(doc, s["commentary"], size=10, space_after=10)

    # Eserciziario
    if eserciziario:
        doc.add_page_break()
        par(doc, "Eserciziario", size=22, bold=True, color=BLU, space_after=4)
        par(doc, "Esercizi presenti solo nel database, non nei file sorgente.",
            size=10, italic=True, color=GRIGIO, space_after=10)
        for e in eserciziario:
            rif = f"esercizio {e.get('id', '')}"
            blocchi = e.get("content") or e.get("blocks")
            if isinstance(blocchi, list):
                for blocco in blocchi:
                    if isinstance(blocco, dict):
                        scrivi_blocco(doc, blocco, rif)
            else:
                par(doc, json.dumps(e, ensure_ascii=False)[:1500], size=9, color=GRIGIO)

    # Glossario
    if glossario:
        doc.add_page_break()
        par(doc, "Glossario", size=22, bold=True, color=BLU, space_after=10)
        for voce in sorted(glossario, key=lambda g: (g.get("term") or "").lower()):
            par(doc, voce.get("term") or "", size=12, bold=True, space_after=1)
            for campo in ("definition", "example", "note"):
                if voce.get(campo):
                    par(doc, str(voce[campo]), size=10, space_after=2,
                        italic=campo != "definition",
                        color=None if campo == "definition" else GRIGIO)
            doc.add_paragraph()

    if anomalie:
        doc.add_page_break()
        par(doc, "Anomalie rilevate automaticamente", size=18, bold=True, color=ROSSO, space_after=6)
        par(doc, "Mani che non contengono 52 carte: vanno corrette prima della "
                 "revisione didattica, sono errori meccanici.", size=10, space_after=8)
        for a in anomalie:
            par(doc, "• " + a, size=10, space_after=2)

    doc.save(OUT)
    print(f"Scritto {OUT}")
    print(f"  {len(corsi)} corsi, {len(lezioni)} lezioni, {len(moduli)} moduli, "
          f"{len(smazzate)} smazzate, {len(glossario)} voci, {len(eserciziario)} esercizi")
    print(f"  anomalie meccaniche: {len(anomalie)}")


if __name__ == "__main__":
    main()
